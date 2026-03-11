"""
Word 文档导出工具
用于将大纲内容导出为 Word 文档
严格遵循文档格式要求.md的格式规范
"""
import io
from typing import List
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 黑色
BLACK = RGBColor(0, 0, 0)


def set_fixed_line_spacing(p, spacing_pts):
    """通过 XML 设置固定值行距"""
    pPr = p._p.get_or_add_pPr()
    if pPr is None:
        return

    # 移除已存在的 spacing 元素
    for child in list(pPr):
        if 'spacing' in child.tag.lower():
            pPr.remove(child)

    # 创建新的 spacing 元素，设置固定值行距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(spacing_pts * 20))  # 28磅 = 560 EMU
    spacing.set(qn('w:lineRule'), 'exact')
    pPr.append(spacing)


def create_plain_paragraph(doc, text, font_name, font_size, bold=False,
                           center=False, first_line_indent=True, heading_level=0):
    """创建纯文本段落，完全自定义格式"""
    p = doc.add_paragraph()
    run = p.add_run(text)

    # 设置字体属性
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = BLACK
    run.font.italic = False
    run.font.underline = False

    # 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 设置对齐
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 设置首行缩进
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    else:
        p.paragraph_format.first_line_indent = None

    # 设置固定值 28 磅行距
    set_fixed_line_spacing(p, 28)

    # 移除边框
    p.paragraph_format.border_bottom = None

    # 设置标题级别
    if heading_level > 0:
        p.style = f'Heading {heading_level}'

    return p


def add_page_break(doc):
    """添加换页符"""
    doc.add_page_break()


def add_header_with_line(section):
    """添加页眉和黑色横线"""
    header = section.header

    # 添加页眉段落
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加下边框（黑色横线）
    pPr = header_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # 线宽
    bottom.set(qn('w:color'), '000000')  # 黑色
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 设置页眉与正文距离
    section.header_distance = Cm(0.3)


def export_outline_to_doc(nodes: List[dict]) -> bytes:
    """
    将大纲树导出为 Word 文档
    严格遵循文档格式要求.md
    """
    doc = Document()

    # 先添加5个空行
    for _ in range(5):
        doc.add_paragraph()

    # 添加主标题 - 黑体 20号 加粗 不倾斜 上下左右居中
    create_plain_paragraph(doc, '人工智能赋能生物育种应用场景项目\n可行性研究报告', '黑体', 20, True, True, False, 0)

    # 添加空行将申报单位推到页面底部（减少空行数量）
    for _ in range(13):
        doc.add_paragraph()

    # 在第一页最下方添加申报单位和日期 - 宋体 四号 不加粗 不倾斜 居中
    create_plain_paragraph(doc, '申报单位：崖州湾国家实验室', '宋体', 14, False, True, False, 0)
    create_plain_paragraph(doc, '2026年01月', '宋体', 14, False, True, False, 0)

    # 换页，开始正文
    add_page_break(doc)

    # 设置第一页不使用页眉（封面页）
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True

    # 从第二章开始，每页添加黑色横线页眉
    # 需要为后续每个 section 添加页眉
    for i, section in enumerate(doc.sections):
        if i > 0:  # 跳过第一页
            add_header_with_line(section)

    def add_node_to_doc(node: dict, doc: Document, is_first=True):
        """递归添加节点"""
        level = node.get('node_level', 1)
        node_code = node.get('node_code', '')
        node_title = node.get('node_title', '')
        content = node.get('content', '')
        children = node.get('children', [])

        # 一级标题前添加换页符，确保新起一页
        if level == 1:
            add_page_break(doc)
            # 每次换页后添加页眉
            # 获取当前 section 并添加页眉
            if len(doc.sections) > 1:
                section = doc.sections[-1]
                add_header_with_line(section)

        # 生成标题文本
        if level == 1:
            heading_text = f"第{node_code}章 {node_title}"
            create_plain_paragraph(doc, heading_text, '黑体', 22, True, True, False, 1)
        else:
            heading_text = f"{node_code} {node_title}"
            heading_level = min(level, 9)
            create_plain_paragraph(doc, heading_text, '宋体', 14, True, False, True, heading_level)

        # 添加内容
        if content and content.strip():
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    create_plain_paragraph(doc, para_text, '宋体', 14, False, False, True, 0)

        # 递归添加子节点
        for child in children:
            add_node_to_doc(child, doc, is_first=False)

    # 遍历所有顶层节点
    for node in nodes:
        add_node_to_doc(node, doc)

    # 保存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
